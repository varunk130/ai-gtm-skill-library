"""
GTM Exec Plan - Sample Word Document Generator
Produces a polished 3-4 page executive GTM brief using Microsoft Fluent design.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Microsoft Fluent colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x00, 0x78, 0xD4)
TEAL = RGBColor(0x00, 0xB2, 0x94)
RED = RGBColor(0xD1, 0x34, 0x38)
GOLD = RGBColor(0xFF, 0xB9, 0x00)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "F3F4F6"
BLUE_HEX = "0078D4"
NAVY_HEX = "1B2A4A"
TEAL_HEX = "00B294"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def style_header_row(row, bg_hex=NAVY_HEX):
    for cell in row.cells:
        set_cell_shading(cell, bg_hex)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Segoe UI Semibold"


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Segoe UI Semibold"
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_header_row(table.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Segoe UI"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Segoe UI Semibold"
        if level == 1:
            run.font.color.rgb = NAVY
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = BLUE
            run.font.size = Pt(13)
        elif level == 3:
            run.font.color.rgb = DARK
            run.font.size = Pt(11)
    return h


def add_body(doc, text, bold=False, color=DARK, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p


def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = "Segoe UI Semibold"
        run_b.font.size = Pt(9.5)
        run_b.font.color.rgb = DARK
        run_b.bold = True
        run_n = p.add_run(text)
        run_n.font.name = "Segoe UI"
        run_n.font.size = Pt(9.5)
        run_n.font.color.rgb = DARK
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def build_cover_page(doc):
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("M365 Copilot Extensibility")
    run.font.name = "Segoe UI Semibold"
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Go-To-Market Strategy")
    run2.font.name = "Segoe UI"
    run2.font.size = Pt(20)
    run2.font.color.rgb = BLUE

    # Tagline
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(20)
    run3 = p3.add_run("Build where your users already work")
    run3.font.name = "Segoe UI"
    run3.font.size = Pt(14)
    run3.font.color.rgb = TEAL
    run3.italic = True

    # Metadata
    for _ in range(4):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = meta.add_run("April 2026  |  Executive Review  |  PRIME Framework")
    run_m.font.name = "Segoe UI"
    run_m.font.size = Pt(10)
    run_m.font.color.rgb = MUTED

    # Disclaimer
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc.paragraph_format.space_before = Pt(12)
    run_d = disc.add_run(
        "NOTE: This is a hypothetical example for demonstration purposes. "
        "All numbers, dates, and targets are illustrative only."
    )
    run_d.font.name = "Segoe UI"
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = MUTED
    run_d.italic = True

    doc.add_page_break()


def build_exec_summary(doc):
    add_heading_styled(doc, "Executive Summary", level=1)
    add_divider(doc)

    add_body(
        doc,
        "Microsoft 365 Copilot Extensibility opens a new developer ecosystem that enables ISVs, "
        "SIs, and enterprise teams to build custom AI agents, plugins, and connectors that run "
        "natively inside M365 Copilot. This represents a platform shift comparable to the early "
        "App Store moment: first movers who build on this platform will capture disproportionate "
        "enterprise distribution.",
        size=10,
    )
    add_body(
        doc,
        "Our GTM motion targets three segments (ISV partners, SI implementers, enterprise dev teams) "
        "with a developer-first acquisition model, partner co-sell amplification, and a 90-day launch "
        "sequence designed to establish category leadership before competitive alternatives mature.",
        size=10,
    )

    # Target callout
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("TARGET: ")
    run.font.name = "Segoe UI Semibold"
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.bold = True
    run2 = p.add_run("500 published extensions and 2,000 active developers within 6 months of GA.")
    run2.font.name = "Segoe UI"
    run2.font.size = Pt(11)
    run2.font.color.rgb = DARK
    run2.bold = True


def build_section_1(doc):
    add_heading_styled(doc, "1. Positioning and Differentiation", level=1)
    add_divider(doc)

    add_heading_styled(doc, "Positioning Statement", level=3)
    add_body(
        doc,
        "For enterprise developers and ISV partners who need to extend AI capabilities within "
        "the tools their users already live in, M365 Copilot Extensibility is the only platform "
        "that provides native integration with the Microsoft 365 graph, security model, and 400M+ "
        "commercial user base, enabling AI-powered workflows that ship in days rather than months.",
        size=10,
    )

    add_heading_styled(doc, "Target Segments", level=3)
    add_styled_table(
        doc,
        ["Segment", "Persona", "Primary Need", "Estimated Size"],
        [
            ["ISV Partners", "Product/Eng Leaders", "Distribution via M365 marketplace", "~8,000 ISVs"],
            ["SI Partners", "Practice Leads", "New billable service line", "~2,500 SI firms"],
            ["Enterprise Dev Teams", "IT/Dev Managers", "Custom agent development", "~45,000 E5 orgs"],
        ],
        col_widths=[1.5, 1.5, 2.0, 1.3],
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Key Differentiators", level=3)
    add_styled_table(
        doc,
        ["vs. Alternative", "Our Advantage"],
        [
            ["Competitor A (AI Assistants)", "Native M365 Graph access, enterprise identity, compliance built in"],
            ["Competitor B (CRM AI)", "Already deployed where users work (Outlook, Teams, Word)"],
            ["Competitor C (Workspace AI)", "3x larger enterprise installed base, deeper LOB app integration"],
            ["Build from scratch", "10x faster to production with pre-built connectors and Copilot Studio"],
        ],
        col_widths=[2.2, 4.2],
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Messaging Hierarchy", level=3)
    add_bullet(doc, " Build where your users already work", bold_prefix="L0 (Tagline):")
    add_bullet(doc, " Ship AI agents that reach 400M+ users through the apps they use every day", bold_prefix="L1 (Value Prop):")
    add_bullet(doc, " Native Graph API, enterprise security, Copilot Studio low-code, marketplace distribution", bold_prefix="L2 (Proof Points):")
    add_bullet(doc, " Declarative agents, API plugins, Graph connectors, Adaptive Cards", bold_prefix="L3 (Technical):")


def build_section_2(doc):
    add_heading_styled(doc, "2. Market Readiness and Competitive Landscape", level=1)
    add_divider(doc)

    add_heading_styled(doc, "Market Opportunity", level=3)
    add_bullet(doc, " $47B enterprise AI platform market (2026, industry estimates)", bold_prefix="TAM:")
    add_bullet(doc, " $12B AI extensibility and integration platform segment", bold_prefix="SAM:")
    add_bullet(doc, " $1.8B addressable through M365 commercial base in Year 1", bold_prefix="SOM:")

    doc.add_paragraph()
    add_heading_styled(doc, "Why Now", level=3)
    add_bullet(doc, "Enterprise AI adoption hit 68% in 2025, but customization remains the #1 unmet need")
    add_bullet(doc, "Competitors are 6-12 months behind on enterprise-grade extensibility frameworks")
    add_bullet(doc, "M365 Copilot seat penetration creates a captive distribution channel")

    doc.add_paragraph()
    add_heading_styled(doc, "Competitive Landscape", level=3)
    add_styled_table(
        doc,
        ["Competitor", "Strength", "Our Advantage", "Threat"],
        [
            ["Competitor A (AI Assistants)", "Brand awareness, developer mindshare", "Enterprise security, M365 Graph, compliance", "High"],
            ["Competitor B (CRM AI)", "CRM-native AI, strong SI channel", "Broader surface area (email, docs, chat)", "Medium"],
            ["Competitor C (IT Workflow AI)", "IT workflow depth", "Platform breadth beyond IT use cases", "Medium"],
            ["Competitor D (Workspace AI)", "Workspace integration", "3x enterprise seat base, LOB app ecosystem", "Medium"],
        ],
        col_widths=[1.8, 1.8, 2.2, 0.7],
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Risk Mitigations", level=3)
    add_styled_table(
        doc,
        ["Risk", "Likelihood", "Mitigation"],
        [
            ["Developer adoption slower than projected", "Medium", "$10M partner incentive fund + Copilot Studio low-code path"],
            ["Competitor A launches enterprise marketplace", "High", "Leverage M365 Graph lock-in and compliance certifications"],
            ["Quality concerns from early extensions", "Medium", "Curated marketplace with certification program at launch"],
        ],
        col_widths=[2.5, 1.0, 3.0],
    )


def build_section_3(doc):
    add_heading_styled(doc, "3. Impact Model and Success Metrics", level=1)
    add_divider(doc)

    add_heading_styled(doc, "Revenue and Adoption Targets", level=3)
    add_styled_table(
        doc,
        ["Milestone", "30 Days", "60 Days", "90 Days", "6 Months"],
        [
            ["Published Extensions", "50", "150", "300", "500"],
            ["Active Developers", "200", "600", "1,200", "2,000"],
            ["Enterprise Orgs Using Extensions", "100", "500", "1,500", "5,000"],
            ["Partner Co-Sell Pipeline", "$5M", "$15M", "$40M", "$120M"],
        ],
        col_widths=[2.2, 1.0, 1.0, 1.0, 1.0],
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Primary KPIs", level=3)
    add_styled_table(
        doc,
        ["Metric", "Target", "Cadence"],
        [
            ["Monthly Active Developers (MAD)", "2,000 by Month 6", "Weekly"],
            ["Extension Publish Rate", "20/week steady state", "Weekly"],
            ["Developer Time-to-First-Extension", "< 4 hours", "Weekly"],
            ["Extension User Retention (D30)", "> 40%", "Monthly"],
            ["Partner Co-Sell Attach Rate", "15% of Copilot deals", "Monthly"],
        ],
        col_widths=[2.8, 2.0, 1.0],
    )


def build_section_4(doc):
    add_heading_styled(doc, "4. Go-To-Market Motion", level=1)
    add_divider(doc)

    add_heading_styled(doc, "Channel Strategy", level=3)
    add_styled_table(
        doc,
        ["Priority", "Channel", "Contribution", "Investment"],
        [
            ["1", "Developer Relations (docs, tutorials, hackathons)", "40%", "$2.5M"],
            ["2", "Partner Co-Sell (ISV + SI activation)", "30%", "$3.0M"],
            ["3", "Digital Marketing (search, social, communities)", "20%", "$1.5M"],
            ["4", "Events (Build, Ignite, partner summits)", "10%", "$1.0M"],
        ],
        col_widths=[0.7, 3.0, 1.1, 1.0],
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Sales Enablement", level=3)
    add_bullet(doc, "Battle cards for top 4 competitive scenarios")
    add_bullet(doc, "15-minute demo script with live extension build")
    add_bullet(doc, "Objection handling guide (security, pricing, migration)")
    add_bullet(doc, "Partner pitch deck for SI/ISV co-sell conversations")

    doc.add_paragraph()
    add_heading_styled(doc, "Partner Ecosystem", level=3)
    add_bullet(doc, " 20 ISV partners building launch-day extensions (confirmed)", bold_prefix="Launch Day:")
    add_bullet(doc, " 5 SI partners offering implementation services at GA", bold_prefix="Services:")
    add_bullet(doc, " 3 enterprise early adopter proof points", bold_prefix="Case Studies:")


def build_section_5(doc):
    add_heading_styled(doc, "5. Execution Timeline", level=1)
    add_divider(doc)

    add_styled_table(
        doc,
        ["Phase", "Timeline", "Key Milestones", "Gate Criteria"],
        [
            ["Pre-Launch", "Weeks -4 to -1", "Messaging final, enablement complete, ISV extensions certified", "Go/No-Go review"],
            ["Launch Week", "Days 1-5", "Press release, CEO demo, developer livestream, SI launch event", "50+ extensions Day 1"],
            ["Post-Launch M1", "Days 8-30", "Weekly metrics cadence, community activation, first hackathon", "200+ MAD, sentiment > 4.0"],
            ["Scale (M2-M3)", "Days 31-90", "Regional expansion, certification program, Phase 2 planning", "1,200 MAD, $40M pipeline"],
        ],
        col_widths=[1.3, 1.2, 2.5, 1.5],
    )

    doc.add_paragraph()
    add_heading_styled(doc, "RACI", level=3)
    add_styled_table(
        doc,
        ["Workstream", "Responsible", "Accountable", "Consulted"],
        [
            ["Messaging and Positioning", "PMM", "CMO", "Product, Sales"],
            ["Developer Experience", "DevRel", "VP Engineering", "PMM, Support"],
            ["Partner Activation", "Partner Team", "CVP Partnerships", "Sales, PMM"],
            ["Sales Enablement", "Enablement", "CRO", "PMM, Product"],
            ["Launch Execution", "GTM Lead", "CMO", "All above"],
        ],
        col_widths=[2.0, 1.5, 1.5, 1.5],
    )


def build_footer(doc):
    doc.add_paragraph()
    add_divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by GTM Exec Plan (PRIME Framework)  |  AI GTM Skill Library")
    run.font.name = "Segoe UI"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Segoe UI"
    font.size = Pt(10)
    font.color.rgb = DARK

    build_cover_page(doc)
    build_exec_summary(doc)
    build_section_1(doc)
    doc.add_page_break()
    build_section_2(doc)
    build_section_3(doc)
    doc.add_page_break()
    build_section_4(doc)
    build_section_5(doc)
    build_footer(doc)

    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "samples")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "gtm-plan-m365-copilot-extensibility.docx")
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
